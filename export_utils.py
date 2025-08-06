import os
import io
from datetime import datetime
from typing import Dict, List, Any
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
import re


class ExportManager:
    """Handles export of AI responses to PDF and Word formats"""
    
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()
    
    def _setup_custom_styles(self):
        """Setup custom styles for better formatting"""
        # Custom styles for PDF - use unique names to avoid conflicts
        self.styles.add(ParagraphStyle(
            name='CustomHeading1',
            parent=self.styles['Heading1'],
            fontSize=16,
            spaceAfter=12,
            spaceBefore=12,
            textColor=colors.darkblue
        ))
        
        self.styles.add(ParagraphStyle(
            name='CustomHeading2',
            parent=self.styles['Heading2'],
            fontSize=14,
            spaceAfter=10,
            spaceBefore=10,
            textColor=colors.darkblue
        ))
        
        self.styles.add(ParagraphStyle(
            name='CustomBodyText',
            parent=self.styles['BodyText'],
            fontSize=11,
            spaceAfter=6,
            alignment=TA_JUSTIFY
        ))
    
    def _clean_html_for_export(self, html_content: str) -> str:
        """Clean HTML content for export, preserving structure"""
        # Remove script and style tags
        html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL)
        html_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL)
        
        # Convert common HTML tags to plain text with structure
        html_content = re.sub(r'<br\s*/?>', '\n', html_content)
        html_content = re.sub(r'<p[^>]*>', '\n', html_content)
        html_content = re.sub(r'</p>', '\n', html_content)
        html_content = re.sub(r'<h1[^>]*>', '\n\n# ', html_content)
        html_content = re.sub(r'</h1>', '\n', html_content)
        html_content = re.sub(r'<h2[^>]*>', '\n\n## ', html_content)
        html_content = re.sub(r'</h2>', '\n', html_content)
        html_content = re.sub(r'<h3[^>]*>', '\n\n### ', html_content)
        html_content = re.sub(r'</h3>', '\n', html_content)
        html_content = re.sub(r'<strong[^>]*>', '**', html_content)
        html_content = re.sub(r'</strong>', '**', html_content)
        html_content = re.sub(r'<em[^>]*>', '*', html_content)
        html_content = re.sub(r'</em>', '*', html_content)
        html_content = re.sub(r'<ul[^>]*>', '\n', html_content)
        html_content = re.sub(r'</ul>', '\n', html_content)
        html_content = re.sub(r'<ol[^>]*>', '\n', html_content)
        html_content = re.sub(r'</ol>', '\n', html_content)
        html_content = re.sub(r'<li[^>]*>', '\n• ', html_content)
        html_content = re.sub(r'</li>', '', html_content)
        
        # Handle tables
        html_content = re.sub(r'<table[^>]*>', '\n\n[TABLE]\n', html_content)
        html_content = re.sub(r'</table>', '\n[/TABLE]\n\n', html_content)
        html_content = re.sub(r'<tr[^>]*>', '', html_content)
        html_content = re.sub(r'</tr>', '\n', html_content)
        html_content = re.sub(r'<th[^>]*>', '| ', html_content)
        html_content = re.sub(r'</th>', ' |', html_content)
        html_content = re.sub(r'<td[^>]*>', '| ', html_content)
        html_content = re.sub(r'</td>', ' |', html_content)
        
        # Remove any remaining HTML tags
        html_content = re.sub(r'<[^>]+>', '', html_content)
        
        # Clean up extra whitespace
        html_content = re.sub(r'\n\s*\n\s*\n', '\n\n', html_content)
        html_content = html_content.strip()
        
        return html_content
    
    def _parse_content_structure(self, content: str) -> List[Dict[str, Any]]:
        """Parse content into structured elements for export"""
        lines = content.split('\n')
        elements = []
        
        for line in lines:
            line = line.strip()
            if not line:
                elements.append({'type': 'spacer'})
                continue
            
            if line.startswith('# '):
                elements.append({
                    'type': 'heading1',
                    'text': line[2:].strip()
                })
            elif line.startswith('## '):
                elements.append({
                    'type': 'heading2',
                    'text': line[3:].strip()
                })
            elif line.startswith('### '):
                elements.append({
                    'type': 'heading3',
                    'text': line[4:].strip()
                })
            elif line.startswith('• '):
                elements.append({
                    'type': 'list_item',
                    'text': line[2:].strip()
                })
            elif line.startswith('[TABLE]'):
                elements.append({'type': 'table_start'})
            elif line.startswith('[/TABLE]'):
                elements.append({'type': 'table_end'})
            elif '|' in line and line.count('|') > 1:
                elements.append({
                    'type': 'table_row',
                    'text': line
                })
            elif line.startswith('**') and line.endswith('**'):
                elements.append({
                    'type': 'bold',
                    'text': line[2:-2].strip()
                })
            elif line.startswith('*') and line.endswith('*'):
                elements.append({
                    'type': 'italic',
                    'text': line[1:-1].strip()
                })
            else:
                elements.append({
                    'type': 'paragraph',
                    'text': line
                })
        
        return elements
    
    def export_to_pdf(self, response_data: Dict[str, Any], filename: str = None) -> bytes:
        """Export AI response to PDF format"""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"ai_response_{timestamp}.pdf"
        
        # Create PDF buffer
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, 
                              topMargin=72, bottomMargin=72)
        
        # Prepare content
        story = []
        
        # Title
        title = Paragraph("AI Response Report", self.styles['CustomHeading1'])
        story.append(title)
        story.append(Spacer(1, 12))
        
        # Timestamp
        timestamp = Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", 
                            self.styles['Normal'])
        story.append(timestamp)
        story.append(Spacer(1, 20))
        
        # Handle single response or multiple responses
        if 'response' in response_data:
            # Single response
            response_text = response_data.get('response', '')
            cleaned_content = self._clean_html_for_export(response_text)
            elements = self._parse_content_structure(cleaned_content)
            
            # Process elements
            for element in elements:
                if element['type'] == 'spacer':
                    story.append(Spacer(1, 6))
                elif element['type'] == 'heading1':
                    story.append(Paragraph(element['text'], self.styles['CustomHeading1']))
                elif element['type'] == 'heading2':
                    story.append(Paragraph(element['text'], self.styles['CustomHeading2']))
                elif element['type'] == 'heading3':
                    story.append(Paragraph(element['text'], self.styles['Heading3']))
                elif element['type'] == 'paragraph':
                    story.append(Paragraph(element['text'], self.styles['CustomBodyText']))
                elif element['type'] == 'bold':
                    story.append(Paragraph(element['text'], self.styles['Heading3']))
                elif element['type'] == 'italic':
                    story.append(Paragraph(element['text'], self.styles['Italic']))
                elif element['type'] == 'list_item':
                    story.append(Paragraph(f"• {element['text']}", self.styles['CustomBodyText']))
                elif element['type'] == 'table_row':
                    # Simple table handling
                    cells = [cell.strip() for cell in element['text'].split('|') if cell.strip()]
                    if cells:
                        table = Table([cells], colWidths=[doc.width/len(cells)]*len(cells))
                        table.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                            ('FONTSIZE', (0, 0), (-1, 0), 12),
                            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                            ('GRID', (0, 0), (-1, -1), 1, colors.black)
                        ]))
                        story.append(table)
                        story.append(Spacer(1, 12))
            
            # Add sources if available
            sources = response_data.get('sources', [])
            if sources:
                story.append(Spacer(1, 20))
                story.append(Paragraph("Sources:", self.styles['CustomHeading2']))
                for i, source in enumerate(sources, 1):
                    story.append(Paragraph(f"{i}. {source}", self.styles['CustomBodyText']))
        
        elif 'responses' in response_data:
            # Multiple responses
            responses = response_data.get('responses', [])
            for i, response_item in enumerate(responses, 1):
                # Add response number
                story.append(Paragraph(f"Response {i}", self.styles['CustomHeading2']))
                story.append(Spacer(1, 10))
                
                # Process response content
                response_text = response_item.get('content', '')
                cleaned_content = self._clean_html_for_export(response_text)
                elements = self._parse_content_structure(cleaned_content)
                
                for element in elements:
                    if element['type'] == 'spacer':
                        story.append(Spacer(1, 6))
                    elif element['type'] == 'heading1':
                        story.append(Paragraph(element['text'], self.styles['CustomHeading1']))
                    elif element['type'] == 'heading2':
                        story.append(Paragraph(element['text'], self.styles['CustomHeading2']))
                    elif element['type'] == 'heading3':
                        story.append(Paragraph(element['text'], self.styles['Heading3']))
                    elif element['type'] == 'paragraph':
                        story.append(Paragraph(element['text'], self.styles['CustomBodyText']))
                    elif element['type'] == 'bold':
                        story.append(Paragraph(element['text'], self.styles['Heading3']))
                    elif element['type'] == 'italic':
                        story.append(Paragraph(element['text'], self.styles['Italic']))
                    elif element['type'] == 'list_item':
                        story.append(Paragraph(f"• {element['text']}", self.styles['CustomBodyText']))
                    elif element['type'] == 'table_row':
                        # Simple table handling
                        cells = [cell.strip() for cell in element['text'].split('|') if cell.strip()]
                        if cells:
                            table = Table([cells], colWidths=[doc.width/len(cells)]*len(cells))
                            table.setStyle(TableStyle([
                                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                                ('FONTSIZE', (0, 0), (-1, 0), 12),
                                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                                ('GRID', (0, 0), (-1, -1), 1, colors.black)
                            ]))
                            story.append(table)
                            story.append(Spacer(1, 12))
                
                # Add sources for this response
                sources = response_item.get('sources', [])
                if sources:
                    story.append(Spacer(1, 10))
                    story.append(Paragraph("Sources:", self.styles['CustomHeading3']))
                    for j, source in enumerate(sources, 1):
                        story.append(Paragraph(f"{j}. {source}", self.styles['CustomBodyText']))
                
                # Add separator between responses
                if i < len(responses):
                    story.append(Spacer(1, 20))
                    story.append(Paragraph("─" * 50, self.styles['Normal']))
                    story.append(Spacer(1, 20))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    def export_to_word(self, response_data: Dict[str, Any], filename: str = None) -> bytes:
        """Export AI response to Word format"""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"ai_response_{timestamp}.docx"
        
        # Create Word document
        doc = Document()
        
        # Title
        title = doc.add_heading('AI Response Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Timestamp
        timestamp_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Add space
        
        # Handle single response or multiple responses
        if 'response' in response_data:
            # Single response
            response_text = response_data.get('response', '')
            cleaned_content = self._clean_html_for_export(response_text)
            elements = self._parse_content_structure(cleaned_content)
            
            # Process elements
            for element in elements:
                if element['type'] == 'spacer':
                    doc.add_paragraph()
                elif element['type'] == 'heading1':
                    doc.add_heading(element['text'], level=1)
                elif element['type'] == 'heading2':
                    doc.add_heading(element['text'], level=2)
                elif element['type'] == 'heading3':
                    doc.add_heading(element['text'], level=3)
                elif element['type'] == 'paragraph':
                    doc.add_paragraph(element['text'])
                elif element['type'] == 'bold':
                    p = doc.add_paragraph()
                    p.add_run(element['text']).bold = True
                elif element['type'] == 'italic':
                    p = doc.add_paragraph()
                    p.add_run(element['text']).italic = True
                elif element['type'] == 'list_item':
                    doc.add_paragraph(element['text'], style='List Bullet')
                elif element['type'] == 'table_row':
                    # Simple table handling
                    cells = [cell.strip() for cell in element['text'].split('|') if cell.strip()]
                    if cells:
                        table = doc.add_table(rows=1, cols=len(cells))
                        table.style = 'Table Grid'
                        for i, cell_text in enumerate(cells):
                            table.cell(0, i).text = cell_text
            
            # Add sources if available
            sources = response_data.get('sources', [])
            if sources:
                doc.add_paragraph()  # Add space
                doc.add_heading('Sources', level=2)
                for i, source in enumerate(sources, 1):
                    doc.add_paragraph(f"{i}. {source}", style='List Number')
        
        elif 'responses' in response_data:
            # Multiple responses
            responses = response_data.get('responses', [])
            for i, response_item in enumerate(responses, 1):
                # Add response number
                doc.add_heading(f"Response {i}", level=2)
                
                # Process response content
                response_text = response_item.get('content', '')
                cleaned_content = self._clean_html_for_export(response_text)
                elements = self._parse_content_structure(cleaned_content)
                
                for element in elements:
                    if element['type'] == 'spacer':
                        doc.add_paragraph()
                    elif element['type'] == 'heading1':
                        doc.add_heading(element['text'], level=1)
                    elif element['type'] == 'heading2':
                        doc.add_heading(element['text'], level=2)
                    elif element['type'] == 'heading3':
                        doc.add_heading(element['text'], level=3)
                    elif element['type'] == 'paragraph':
                        doc.add_paragraph(element['text'])
                    elif element['type'] == 'bold':
                        p = doc.add_paragraph()
                        p.add_run(element['text']).bold = True
                    elif element['type'] == 'italic':
                        p = doc.add_paragraph()
                        p.add_run(element['text']).italic = True
                    elif element['type'] == 'list_item':
                        doc.add_paragraph(element['text'], style='List Bullet')
                    elif element['type'] == 'table_row':
                        # Simple table handling
                        cells = [cell.strip() for cell in element['text'].split('|') if cell.strip()]
                        if cells:
                            table = doc.add_table(rows=1, cols=len(cells))
                            table.style = 'Table Grid'
                            for i, cell_text in enumerate(cells):
                                table.cell(0, i).text = cell_text
                
                # Add sources for this response
                sources = response_item.get('sources', [])
                if sources:
                    doc.add_paragraph()  # Add space
                    doc.add_heading('Sources', level=3)
                    for j, source in enumerate(sources, 1):
                        doc.add_paragraph(f"{j}. {source}", style='List Number')
                
                # Add separator between responses
                if i < len(responses):
                    doc.add_paragraph("─" * 50)
                    doc.add_paragraph()
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def get_export_filename(self, format_type: str, base_name: str = "ai_response") -> str:
        """Generate a filename for export"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        if format_type.lower() == 'pdf':
            return f"{base_name}_{timestamp}.pdf"
        elif format_type.lower() == 'docx':
            return f"{base_name}_{timestamp}.docx"
        else:
            return f"{base_name}_{timestamp}.{format_type}" 